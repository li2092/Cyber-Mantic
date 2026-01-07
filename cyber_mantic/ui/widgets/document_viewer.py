"""
文档查看器组件
支持 Markdown / 文本 / Word / PDF 文件的阅读

功能：
- Markdown 渲染
- Word 文档渲染（需要 python-docx）
- PDF 文档渲染（需要 PyMuPDF）
- 文本选择
- 选中文本工具栏

设计参考：docs/design/02_典籍模块设计.md
"""
import re
from pathlib import Path
from typing import Optional, Callable
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QTextBrowser, QLabel,
    QToolBar, QPushButton, QMenu, QHBoxLayout, QFrame
)
from PyQt6.QtCore import Qt, pyqtSignal, QPoint
from PyQt6.QtGui import QAction, QFont, QTextCursor

from utils.logger import get_logger

# 可选依赖
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False


class DocumentViewer(QWidget):
    """
    文档查看器

    支持 Markdown 文件的渲染和阅读
    """

    # 信号：选中文本请求创建笔记
    note_requested = pyqtSignal(str, str, str)  # (content, source_file, position)
    # 信号：选中文本请求提问
    question_requested = pyqtSignal(str)  # (selected_text)
    # 信号：滚动检测（用于阅读计时）
    scroll_detected = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.logger = get_logger(__name__)

        self._current_file: Optional[str] = None
        self._current_title: str = ""

        self._init_ui()

    def _init_ui(self):
        """初始化UI"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(5)

        # 文档标题栏
        self._title_bar = self._create_title_bar()
        layout.addWidget(self._title_bar)

        # 文档内容区
        self._browser = QTextBrowser()
        self._browser.setOpenExternalLinks(False)
        self._browser.setStyleSheet("""
            QTextBrowser {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 15px;
                background-color: #fff;
                font-size: 14px;
                line-height: 1.6;
            }
        """)

        # 连接选择变化信号
        self._browser.selectionChanged.connect(self._on_selection_changed)

        # 连接滚动信号（用于阅读计时）
        self._browser.verticalScrollBar().valueChanged.connect(self._on_scroll)

        layout.addWidget(self._browser)

        # 选中文本工具栏（初始隐藏）
        self._selection_toolbar = self._create_selection_toolbar()
        self._selection_toolbar.setVisible(False)
        layout.addWidget(self._selection_toolbar)

    def _create_title_bar(self) -> QWidget:
        """创建标题栏"""
        bar = QFrame()
        bar.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 5px;
                padding: 5px;
            }
        """)

        layout = QHBoxLayout(bar)
        layout.setContentsMargins(10, 5, 10, 5)

        self._title_label = QLabel("请选择文档")
        self._title_label.setStyleSheet("font-size: 14px; font-weight: bold; color: #333;")
        layout.addWidget(self._title_label)

        layout.addStretch()

        self._info_label = QLabel("")
        self._info_label.setStyleSheet("font-size: 12px; color: #666;")
        layout.addWidget(self._info_label)

        return bar

    def _create_selection_toolbar(self) -> QWidget:
        """创建选中文本工具栏"""
        toolbar = QFrame()
        toolbar.setStyleSheet("""
            QFrame {
                background-color: #e3f2fd;
                border: 1px solid #90caf9;
                border-radius: 5px;
                padding: 5px;
            }
        """)

        layout = QHBoxLayout(toolbar)
        layout.setContentsMargins(10, 5, 10, 5)

        hint_label = QLabel("已选中文本：")
        hint_label.setStyleSheet("color: #1976d2;")
        layout.addWidget(hint_label)

        # 添加笔记按钮
        note_btn = QPushButton("📝 添加笔记")
        note_btn.clicked.connect(self._on_add_note)
        note_btn.setStyleSheet("""
            QPushButton {
                padding: 5px 15px;
                border: 1px solid #1976d2;
                border-radius: 3px;
                background-color: #fff;
                color: #1976d2;
            }
            QPushButton:hover {
                background-color: #e3f2fd;
            }
        """)
        layout.addWidget(note_btn)

        # 复制按钮
        copy_btn = QPushButton("📋 复制")
        copy_btn.clicked.connect(self._on_copy)
        copy_btn.setStyleSheet("""
            QPushButton {
                padding: 5px 15px;
                border: 1px solid #ccc;
                border-radius: 3px;
                background-color: #fff;
            }
            QPushButton:hover {
                background-color: #f5f5f5;
            }
        """)
        layout.addWidget(copy_btn)

        # 提问按钮
        question_btn = QPushButton("💬 提问")
        question_btn.clicked.connect(self._on_question)
        question_btn.setStyleSheet("""
            QPushButton {
                padding: 5px 15px;
                border: 1px solid #ccc;
                border-radius: 3px;
                background-color: #fff;
            }
            QPushButton:hover {
                background-color: #f5f5f5;
            }
        """)
        layout.addWidget(question_btn)

        layout.addStretch()

        # 取消选择按钮
        cancel_btn = QPushButton("✕")
        cancel_btn.clicked.connect(self._clear_selection)
        cancel_btn.setStyleSheet("""
            QPushButton {
                padding: 5px 10px;
                border: none;
                color: #666;
            }
            QPushButton:hover {
                color: #333;
            }
        """)
        layout.addWidget(cancel_btn)

        return toolbar

    def load_file(self, file_path: str) -> bool:
        """
        加载文件

        Args:
            file_path: 文件路径

        Returns:
            是否成功
        """
        path = Path(file_path)

        if not path.exists():
            self.logger.error(f"文件不存在: {file_path}")
            self._show_error(f"文件不存在: {path.name}")
            return False

        try:
            suffix = path.suffix.lower()

            if suffix == '.md':
                content = path.read_text(encoding='utf-8')
                html = self._markdown_to_html(content)
                self._browser.setHtml(html)

            elif suffix == '.txt':
                content = path.read_text(encoding='utf-8')
                self._browser.setPlainText(content)

            elif suffix in ('.doc', '.docx'):
                if HAS_DOCX:
                    html = self._docx_to_html(path)
                    if html:
                        self._browser.setHtml(html)
                    else:
                        self._show_error("Word文档解析失败")
                        return False
                else:
                    self._show_placeholder("需要安装 python-docx 库来阅读 Word 文档\n\npip install python-docx")
                    return False

            elif suffix == '.pdf':
                if HAS_PYMUPDF:
                    html = self._pdf_to_html(path)
                    if html:
                        self._browser.setHtml(html)
                    else:
                        self._show_error("PDF文档解析失败")
                        return False
                else:
                    self._show_placeholder("需要安装 PyMuPDF 库来阅读 PDF 文档\n\npip install pymupdf")
                    return False

            else:
                # 尝试作为文本读取
                try:
                    content = path.read_text(encoding='utf-8')
                    self._browser.setPlainText(content)
                except Exception:
                    self._show_error(f"不支持的文件格式: {suffix}")
                    return False

            self._current_file = str(path)
            self._current_title = path.stem
            self._title_label.setText(self._current_title)
            self._info_label.setText(f"格式: {suffix[1:].upper()}")

            self.logger.info(f"加载文档: {file_path}")
            return True

        except Exception as e:
            self.logger.error(f"加载文件失败: {e}")
            self._show_error(f"加载失败: {str(e)}")
            return False

    def _markdown_to_html(self, markdown_text: str) -> str:
        """
        将 Markdown 转换为 HTML

        简单实现，支持基本语法
        """
        html = markdown_text

        # 转义 HTML 特殊字符（保留 Markdown 标记）
        html = html.replace('&', '&amp;')
        html = html.replace('<', '&lt;')
        html = html.replace('>', '&gt;')

        # 恢复 Markdown 使用的 > 符号（引用块）
        lines = html.split('\n')
        processed_lines = []
        in_code_block = False

        for line in lines:
            # 代码块
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                if in_code_block:
                    processed_lines.append('<pre style="background-color: #f5f5f5; padding: 10px; border-radius: 5px; overflow-x: auto;">')
                else:
                    processed_lines.append('</pre>')
                continue

            if in_code_block:
                processed_lines.append(line)
                continue

            # 标题
            if line.startswith('# '):
                line = f'<h1 style="border-bottom: 1px solid #eee; padding-bottom: 10px;">{line[2:]}</h1>'
            elif line.startswith('## '):
                line = f'<h2 style="border-bottom: 1px solid #eee; padding-bottom: 8px;">{line[3:]}</h2>'
            elif line.startswith('### '):
                line = f'<h3>{line[4:]}</h3>'
            elif line.startswith('#### '):
                line = f'<h4>{line[5:]}</h4>'

            # 引用块
            elif line.startswith('&gt; '):
                line = f'<blockquote style="border-left: 4px solid #ddd; padding-left: 15px; color: #666; margin: 10px 0;">{line[5:]}</blockquote>'

            # 无序列表
            elif line.strip().startswith('- '):
                line = f'<li>{line.strip()[2:]}</li>'
            elif line.strip().startswith('* '):
                line = f'<li>{line.strip()[2:]}</li>'

            # 有序列表
            elif re.match(r'^\d+\. ', line.strip()):
                content = re.sub(r'^\d+\. ', '', line.strip())
                line = f'<li>{content}</li>'

            # 分隔线
            elif line.strip() in ('---', '***', '___'):
                line = '<hr style="border: none; border-top: 1px solid #ddd; margin: 20px 0;">'

            # 空行
            elif line.strip() == '':
                line = '<br>'

            else:
                # 普通段落
                line = f'<p style="margin: 10px 0; line-height: 1.8;">{line}</p>'

            processed_lines.append(line)

        html = '\n'.join(processed_lines)

        # 行内样式
        # 粗体
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'__(.+?)__', r'<strong>\1</strong>', html)

        # 斜体
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        html = re.sub(r'_(.+?)_', r'<em>\1</em>', html)

        # 行内代码
        html = re.sub(r'`(.+?)`', r'<code style="background-color: #f5f5f5; padding: 2px 5px; border-radius: 3px;">\1</code>', html)

        # 链接
        html = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2" style="color: #1976d2;">\1</a>', html)

        # 包装为完整 HTML
        return self._wrap_html(html)

    def _docx_to_html(self, file_path: Path) -> Optional[str]:
        """
        将 Word 文档转换为 HTML

        Args:
            file_path: 文件路径

        Returns:
            HTML字符串，失败返回None
        """
        try:
            doc = DocxDocument(str(file_path))
            html_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    html_parts.append('<br>')
                    continue

                # 检查段落样式
                style_name = para.style.name if para.style else ""

                if 'Heading 1' in style_name or style_name == '标题 1':
                    html_parts.append(f'<h1>{text}</h1>')
                elif 'Heading 2' in style_name or style_name == '标题 2':
                    html_parts.append(f'<h2>{text}</h2>')
                elif 'Heading 3' in style_name or style_name == '标题 3':
                    html_parts.append(f'<h3>{text}</h3>')
                elif 'Heading' in style_name:
                    html_parts.append(f'<h4>{text}</h4>')
                else:
                    # 处理段落中的格式
                    formatted_text = self._format_docx_runs(para)
                    html_parts.append(f'<p style="margin: 10px 0; line-height: 1.8;">{formatted_text}</p>')

            # 处理表格
            for table in doc.tables:
                html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 15px 0;">')
                for row in table.rows:
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        html_parts.append(f'<td style="border: 1px solid #ddd; padding: 8px;">{cell.text}</td>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')

            return self._wrap_html('\n'.join(html_parts))

        except Exception as e:
            self.logger.error(f"解析Word文档失败: {e}")
            return None

    def _format_docx_runs(self, para) -> str:
        """格式化段落中的文本run"""
        result = []
        for run in para.runs:
            text = run.text
            if not text:
                continue

            # 应用格式
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'

            result.append(text)

        return ''.join(result) if result else para.text

    def _pdf_to_html(self, file_path: Path) -> Optional[str]:
        """
        将 PDF 文档转换为 HTML

        Args:
            file_path: 文件路径

        Returns:
            HTML字符串，失败返回None
        """
        try:
            doc = fitz.open(str(file_path))
            html_parts = []

            html_parts.append(f'<div style="text-align: center; color: #666; margin-bottom: 20px;">')
            html_parts.append(f'共 {len(doc)} 页')
            html_parts.append('</div>')

            for page_num, page in enumerate(doc, 1):
                # 添加页面分隔
                if page_num > 1:
                    html_parts.append('<hr style="border: none; border-top: 2px dashed #ddd; margin: 30px 0;">')

                html_parts.append(f'<div style="color: #999; font-size: 12px; margin-bottom: 10px;">第 {page_num} 页</div>')

                # 提取文本
                text = page.get_text()
                if text.strip():
                    # 按段落分割
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            # 转义HTML
                            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            para = para.replace('\n', '<br>')
                            html_parts.append(f'<p style="margin: 10px 0; line-height: 1.8;">{para}</p>')
                else:
                    html_parts.append('<p style="color: #999;">[此页无文本内容，可能是图片或扫描页]</p>')

            doc.close()
            return self._wrap_html('\n'.join(html_parts))

        except Exception as e:
            self.logger.error(f"解析PDF文档失败: {e}")
            return None

    def _wrap_html(self, content: str) -> str:
        """包装HTML内容"""
        return f'''
        <html>
        <head>
            <style>
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                h1, h2, h3, h4 {{ margin-top: 20px; margin-bottom: 10px; }}
                ul, ol {{ padding-left: 30px; }}
                li {{ margin: 5px 0; }}
                table {{ border-collapse: collapse; }}
            </style>
        </head>
        <body>
            {content}
        </body>
        </html>
        '''

    def _show_placeholder(self, message: str):
        """显示占位信息"""
        self._browser.setHtml(f'''
            <div style="text-align: center; padding: 100px 20px; color: #666;">
                <p style="font-size: 48px; margin-bottom: 20px;">🚧</p>
                <p style="font-size: 18px;">{message}</p>
            </div>
        ''')

    def _show_error(self, message: str):
        """显示错误信息"""
        self._browser.setHtml(f'''
            <div style="text-align: center; padding: 100px 20px; color: #d32f2f;">
                <p style="font-size: 48px; margin-bottom: 20px;">❌</p>
                <p style="font-size: 18px;">{message}</p>
            </div>
        ''')

    def show_welcome(self):
        """显示欢迎页面"""
        self._browser.setHtml('''
            <div style="text-align: center; padding: 80px 20px; color: #666;">
                <p style="font-size: 48px; margin-bottom: 20px;">📚</p>
                <h2 style="color: #333; margin-bottom: 20px;">典籍阅读</h2>
                <p style="font-size: 16px; margin-bottom: 30px;">
                    从左侧选择文档开始阅读
                </p>
                <div style="text-align: left; max-width: 400px; margin: 0 auto; background: #f9f9f9; padding: 20px; border-radius: 10px;">
                    <p style="margin-bottom: 10px;"><strong>功能提示：</strong></p>
                    <ul style="line-height: 2; color: #555;">
                        <li>选中文本可添加笔记</li>
                        <li>支持 Markdown 格式文档</li>
                        <li>可使用 AI 助手辅助学习</li>
                    </ul>
                </div>
            </div>
        ''')
        self._title_label.setText("请选择文档")
        self._info_label.setText("")
        self._current_file = None

    def _on_scroll(self):
        """滚动事件处理"""
        self.scroll_detected.emit()

    def _on_selection_changed(self):
        """选择变化时显示/隐藏工具栏"""
        cursor = self._browser.textCursor()
        has_selection = cursor.hasSelection()
        self._selection_toolbar.setVisible(has_selection)

    def _get_selected_text(self) -> str:
        """获取选中的文本"""
        cursor = self._browser.textCursor()
        return cursor.selectedText().replace('\u2029', '\n')

    def _on_add_note(self):
        """添加笔记"""
        selected = self._get_selected_text()
        if selected:
            self.note_requested.emit(
                selected,
                self._current_file or "",
                ""  # position - 可以后续增强
            )

    def _on_copy(self):
        """复制选中文本"""
        self._browser.copy()

    def _on_question(self):
        """提问"""
        selected = self._get_selected_text()
        if selected:
            self.question_requested.emit(selected)

    def _clear_selection(self):
        """清除选择"""
        cursor = self._browser.textCursor()
        cursor.clearSelection()
        self._browser.setTextCursor(cursor)
        self._selection_toolbar.setVisible(False)

    @property
    def current_file(self) -> Optional[str]:
        """当前打开的文件"""
        return self._current_file

    @property
    def current_title(self) -> str:
        """当前文档标题"""
        return self._current_title
